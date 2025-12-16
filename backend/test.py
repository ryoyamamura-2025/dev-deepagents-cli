# import argparse
# from PIL import Image
# import os
# import sys
# from pathlib import Path

# from google import genai
# from google.genai import types

# # Default location for Vertex AI (global for image generation models)
# DEFAULT_LOCATION = "global"


# def generate_image(content: str, output_dir: str, style: str = "detail", length: str = "5~", lang: str = "jp") -> None:
#     """Generate image using Gemini API and save to file."""

#     full_system_prompt = f"どんなリクエストに対しても画像を2枚以上生成して"

#     try:
#         # Use ADC (Application Default Credentials) via Vertex AI
#         project = os.environ.get("GOOGLE_CLOUD_PROJECT")  # None -> auto-detect from ADC
#         location = DEFAULT_LOCATION
#         client = genai.Client(
#             vertexai=True,
#             project=project,
#             location=location
#         )

#         response = client.models.generate_content(
#             model="gemini-3-pro-image-preview",  # Nano Banana Pro
#             contents=content,
#             config=types.GenerateContentConfig(
#                 system_instruction=[full_system_prompt],
#                 response_modalities=["image", "text"],
#                 image_config=types.ImageConfig(
#                     aspect_ratio="16:9",
#                     image_size="1K"
#                 )
#             )
#         )

#         # Extract image from response
#         i = 0
#         for part in response.parts:
#             if part.text is not None:
#                 print(f"Response text: {part.text}")
#             elif part.inline_data is not None:
#                 image = part.as_image()
#                 image.save(f"{output_dir}/slide_{i}.png")
#                 print(f"✅ Generated: {output_dir}/slide_{i}.png")
#                 i += 1
#         if i == 0:
#             print("Error: No image was generated in the response", file=sys.stderr)
#             sys.exit(1)
#         else:
#             return

#     except Exception as e:
#         print(f"Error generating image: {e}", file=sys.stderr)
#         sys.exit(1)


# def main():
#     parser = argparse.ArgumentParser(
#         description="Generate images using Gemini API",
#         formatter_class=argparse.RawDescriptionHelpFormatter,
#         epilog=__doc__
#     )
#     parser.add_argument(
#         "output_dir",
#         help="Output file directory (e.g., WORKING_DIR/flow/slide-generator/datetime)"
#     )
#     parser.add_argument(
#         "--style",
#         choices=["detail", "presentor"],
#         default="detail",
#         help="Slide Content volume (default: detail)"
#     )
#     parser.add_argument(
#         "--length",
#         choices=["1", "3", "5~"],
#         default="5~",
#         help="Number of slides (default: 5~)"
#     )
#     parser.add_argument(
#         "--lang",
#         choices=["jp", "en"],
#         default="jp",
#         help="Language (default: jp)"
#     )

#     args = parser.parse_args()

#     content = "かわいい猫"

#     generate_image(content, args.output_dir, args.style, args.length, args.lang)


# if __name__ == "__main__":
#     main()

import json
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

# ==========================================
# 設定・定数
# ==========================================
THEME_COLOR = RGBColor(102, 51, 204)  # PDFの紫に近い色 (Dark Slate Blue系)
LIGHT_BG_COLOR = "E6E6FA"             # 概要欄の薄い背景色 (Lavender)
TAG_BG_COLOR = "6633CC"               # タグの背景色

# ==========================================
# ヘルパー関数 (OXML操作による装飾用)
# ==========================================

def set_cell_bg_color(cell, color_hex):
    """表のセルに背景色を設定する"""
    cell_properties = cell._element.tcPr
    try:
        cell_shading = cell_properties.xpath('w:shd')[0]
    except IndexError:
        cell_shading = OxmlElement('w:shd')
    cell_shading.set(qn('w:fill'), color_hex)
    cell_properties.append(cell_shading)

def add_paragraph_left_border(paragraph, color_hex="6633CC", sz="24", space="10"):
    """段落の左側に太い縦線を追加する"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)        # 線の太さ (1/8 pt単位) -> 24 = 3pt
    left.set(qn('w:space'), space)  # テキストとの距離
    left.set(qn('w:color'), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)

def set_text_shading(run, color_hex):
    """テキストランの背景色（ハイライトではない網掛け）を設定する"""
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    rPr.append(shd)

def create_dummy_image(filename):
    """
    【テスト用】画像ファイルがない場合にダミー画像を作成する。
    本番環境で画像が揃っている場合は不要ですが、エラー回避のため入れています。
    """
    from PIL import Image, ImageDraw, ImageFont
    if not os.path.exists(filename):
        img = Image.new('RGB', (600, 338), color=(240, 240, 240))
        d = ImageDraw.Draw(img)
        d.text((10, 150), f"Image: {filename}", fill=(0, 0, 0))
        img.save(filename)

# ==========================================
# メイン処理
# ==========================================

def create_manual_docx(json_data, output_filename="manual.docx"):
    doc = Document()

    # スタイル調整（デフォルトフォントなど）
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Meiryo UI' # Windows向けフォント設定
    font.size = Pt(10.5)

    # 1. マニュアルタイトル
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(json_data['manual_title'])
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 50) # 濃い紺色
    title_p.space_after = Pt(12)

    # 2. 概要（Abstract） - 薄紫のボックス
    # 1行1列の表を作成して背景色をつける
    abstract_table = doc.add_table(rows=1, cols=1)
    abstract_table.autofit = False
    abstract_table.columns[0].width = Inches(6.5) # A4幅に合わせて調整
    
    cell = abstract_table.cell(0, 0)
    set_cell_bg_color(cell, LIGHT_BG_COLOR)
    
    # 概要テキスト
    abs_p = cell.paragraphs[0]
    abs_p.add_run(json_data['manual_abstract'])
    # 余白調整
    abs_p.paragraph_format.left_indent = Pt(5)
    abs_p.paragraph_format.right_indent = Pt(5)
    abs_p.paragraph_format.space_before = Pt(5)
    abs_p.paragraph_format.space_after = Pt(5)

    doc.add_paragraph() # スペーサー

    # 3. セクションのループ
    for section in json_data['sections']:
        # セクションタイトル
        h1 = doc.add_heading(section['section_title'], level=1)
        # 見出しの色を調整
        for run in h1.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # セクション下のライン（Wordの見出しスタイル依存だが、ここでは明示的に下線を引く代わりにそのまま）
        
        # 4. 行（Row）のループ
        for row in section['rows']:
            # レイアウト用の表を作成 (左: 画像, 右: テキスト)
            # 枠線なしの表
            row_table = doc.add_table(rows=1, cols=2)
            row_table.autofit = False
            
            # 列幅の調整 (左: 60%, 右: 40% くらい)
            row_table.columns[0].width = Inches(3.8)
            row_table.columns[1].width = Inches(2.7)

            # セルを取得
            left_cell = row_table.cell(0, 0)
            right_cell = row_table.cell(0, 1)
            
            # --- 左カラム：画像 ---
            left_p = left_cell.paragraphs[0]
            left_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for shot in row['screenshots']:
                ts = shot['timestamp']
                # ファイル名を生成 (例: screenshot_00:33.233.png -> screenshot_00_33_233.png など、ファイルシステムに合わせて調整)
                # ここではJSONのタイムスタンプそのままの名前を想定し、無効文字(:)の置換処理を入れる
                safe_ts = ts.replace(':', '.')
                img_filename = f"screenshot_{safe_ts}.png"
                
                # ダミー画像生成（テスト用）
                create_dummy_image(img_filename)
                
                # 画像貼り付け
                try:
                    run = left_p.add_run()
                    run.add_picture(img_filename, width=Inches(3.6))
                    # キャプション
                    cap_run = left_p.add_run(f"\n{shot['caption']}\n")
                    cap_run.font.size = Pt(9)
                    cap_run.font.color.rgb = RGBColor(80, 80, 80)
                except Exception as e:
                    left_p.add_run(f"[画像読み込みエラー: {img_filename}]")

            # --- 右カラム：テキスト情報 ---
            right_p = right_cell.paragraphs[0]
            
            # (1) タグとタイムスタンプ
            # タグ (Operation, Checkなど)
            tag_run = right_p.add_run(f" {row['tag']} ")
            tag_run.font.color.rgb = RGBColor(255, 255, 255) # 白文字
            tag_run.font.size = Pt(9)
            set_text_shading(tag_run, TAG_BG_COLOR) # 紫背景
            
            # タイムスタンプ
            time_run = right_p.add_run(f"\t{row['timestamp_start']}")
            time_run.font.size = Pt(11)
            time_run.font.bold = True
            
            # (2) 見出し（タイトル）
            # 新しい段落を追加
            heading_p = right_cell.add_paragraph()
            heading_p.paragraph_format.space_before = Pt(6)
            heading_p.paragraph_format.space_after = Pt(6)
            
            # 左側に紫のラインを追加する関数呼び出し
            add_paragraph_left_border(heading_p, color_hex=TAG_BG_COLOR, sz="32", space="120")
            
            head_run = heading_p.add_run(row['heading'])
            head_run.font.size = Pt(14)
            head_run.font.bold = True
            head_run.font.color.rgb = THEME_COLOR

            # (3) メインテキスト
            content_p = right_cell.add_paragraph(row['content']['main_text'])
            content_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            content_p.paragraph_format.line_spacing = 1.2
            content_p.paragraph_format.space_after = Pt(8)

            # (4) 箇条書き
            for point in row['content']['bullet_points']:
                bul_p = right_cell.add_paragraph()
                # PDFのような四角い記号(☐)にするか、通常の箇条書きにするか
                # ここでは文字としての四角を使用
                bul_run = bul_p.add_run("☐ " + point)
                bul_p.paragraph_format.left_indent = Pt(0) # インデント調整
                bul_p.paragraph_format.space_after = Pt(2)
            
            # テーブルの後に余白を入れる
            doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 保存
    doc.save(output_filename)
    print(f"Wordファイルを作成しました: {output_filename}")

# ==========================================
# 実行部分
# ==========================================

# 質問のJSONデータ
input_json = {
 "manual_title": "タイミングチャートを用いたJAM 4711の解析手順",
 "manual_abstract": "このマニュアルでは、タイミングチャートを用いてJAM 4711の発生原因を特定する解析手順を解説します。センサーのオン・オフタイミングを比較し、用紙の搬送状態を分析することで、問題の根本原因を明らかにします。",
 "sections": [
  {
   "section_title": "JAMの発生状況確認",
   "rows": [
    {
     "timestamp_start": "00:17",
     "heading": "JAM 4711の発生条件確認",
     "tag": "Check",
     "content": {
      "main_text": "JAM 4711は「排出反転センサーがオフしない」という条件で発生します。タイミングチャートで、ジャム発生時に排出反転センサーがオンのままであることを確認します。",
      "bullet_points": [
       "排出反転センサーの波形を確認します。",
       "センサーがオンになった後、ジャム検知（13.97s）までオフになっていないことがわかります。"
      ]
     },
     "screenshots": [
      {
       "timestamp": "00:33.233",
       "caption": "排出反転センサーの波形を確認します。"
      },
      {
       "timestamp": "00:47.447",
       "caption": "ジャム検知までセンサーがオフになっていないことを確認します。"
      }
     ]
    }
   ]
  },
  {
   "section_title": "原因の特定と結論",
   "rows": [
    {
     "timestamp_start": "01:15",
     "heading": "用紙先頭の通過タイミング比較",
     "tag": "Analysis",
     "content": {
      "main_text": "ジャムした用紙の搬送が正常だったかを確認するため、正常な用紙（1枚前、2枚前）と搬送タイミングを比較します。まず、用紙の先頭がセンサーを通過するタイミング（センサーON）を比較します。",
      "bullet_points": [
       "「定着排出センサー」のONから「排出反転センサー」のONまでの時間を計測します。",
       "2枚前の用紙：390ms",
       "1枚前の用紙：382ms",
       "ジャムした用紙：388ms",
       "結果として、用紙先頭の通過タイミングに大きな差は見られません。"
      ]
     },
     "screenshots": [
      {
       "timestamp": "02:15.815",
       "caption": "2枚前の用紙のセンサー間通過時間を計測します。"
      },
      {
       "timestamp": "02:25.825",
       "caption": "1枚前の用紙のセンサー間通過時間を計測します。"
      },
      {
       "timestamp": "03:07.307",
       "caption": "ジャムした用紙のセンサー間通過時間を計測します。"
      }
     ]
    },
    {
     "timestamp_start": "03:45",
     "heading": "用紙後端の通過タイミング比較",
     "tag": "Analysis",
     "content": {
      "main_text": "次に、用紙の後端がセンサーを通過するタイミングを比較します。ここでは「排出反転センサー」のONから「両面搬送センサー1」のONまでの時間を比較します。",
      "bullet_points": [
       "2枚前の用紙：656ms",
       "1枚前の用紙：652ms",
       "ジャムした用紙：676ms",
       "結果として、ジャムした用紙は約20ms遅れて搬送されていることがわかります。"
      ]
     },
     "screenshots": [
      {
       "timestamp": "04:44.544",
       "caption": "2枚前の用紙の搬送時間を計測します。"
      },
      {
       "timestamp": "04:25.425",
       "caption": "1枚前の用紙の搬送時間を計測します。"
      },
      {
       "timestamp": "04:37.437",
       "caption": "ジャムした用紙の搬送時間を計測します。"
      }
     ]
    },
    {
     "timestamp_start": "05:33",
     "heading": "原因の推定と用紙長の確認",
     "tag": "Analysis",
     "content": {
      "main_text": "搬送遅延の原因として「重送」が考えられます。重送が発生しているかを確認するため、用紙長に相当する「定着排出センサー」のオン時間を計測・比較します。",
      "bullet_points": [
       "2枚前の用紙：532ms",
       "1枚前の用紙：530ms",
       "ジャムした用紙：556ms",
       "結果として、ジャムした用紙のオン時間が約20ms長くなっています。これは用紙長が長い、つまり重送が発生していることを示唆します。"
      ]
     },
     "screenshots": [
      {
       "timestamp": "08:00.800",
       "caption": "2枚前の用紙の用紙長（センサーON時間）を計測します。"
      },
      {
       "timestamp": "08:07.307",
       "caption": "1枚前の用紙の用紙長を計測します。"
      },
      {
       "timestamp": "08:12.312",
       "caption": "ジャムした用紙の用紙長を計測します。"
      }
     ]
    },
    {
     "timestamp_start": "09:24",
     "heading": "結論",
     "tag": "Conclusion",
     "content": {
      "main_text": "以上の解析から、JAM 4711の発生原因は重送であると結論付けられます。",
      "bullet_points": [
       "重送により、ジャムした用紙の検知上の用紙長が通常より長くなりました。",
       "これにより搬送遅延が発生し、排出反転位置で用紙が分離しました。",
       "分離した後続の用紙が排出反転センサー上に残り続けました。",
       "結果として、排出反転センサーがオフにならず、JAM 4711が検知されました。"
      ]
     },
     "screenshots": [
      {
       "timestamp": "09:36.536",
       "caption": "解析結果のまとめです。"
      }
     ]
    }
   ]
  }
 ]
}

# 実行
if __name__ == "__main__":
    create_manual_docx(input_json)